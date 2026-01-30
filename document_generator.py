"""
ドキュメント生成モジュール - Word/PDF出力
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
import tempfile
import os
import logging
import re
import glob
from typing import Dict
from datetime import datetime

logger = logging.getLogger(__name__)

class JapanesePDF(FPDF):
    """日本語対応PDF生成クラス"""

    def __init__(self):
        super().__init__()
        self.font_name = None
        self._setup_japanese_font()

    def _setup_japanese_font(self):
        """日本語フォントを設定"""
        # 日本語フォントパス（各OS用）
        font_paths = [
            # Windows
            ("C:\\Windows\\Fonts\\msgothic.ttc", "MSGothic"),
            ("C:\\Windows\\Fonts\\meiryo.ttc", "Meiryo"),
            ("C:\\Windows\\Fonts\\YuGothR.ttc", "YuGothic"),
            ("C:\\Windows\\Fonts\\msmincho.ttc", "MSMincho"),
            # macOS
            ("/System/Library/Fonts/ヒラギノ角ゴシック W3.ttc", "Hiragino"),
            ("/Library/Fonts/Arial Unicode.ttf", "ArialUnicode"),
            # Linux (Debian/Ubuntu fonts-noto-cjk) - 優先順位順
            ("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc", "NotoSansCJK"),
            ("/usr/share/fonts/opentype/noto/NotoSansCJKjp-Regular.otf", "NotoSansCJK"),
            ("/usr/share/fonts/truetype/noto/NotoSansCJKjp-Regular.otf", "NotoSansCJK"),
            ("/usr/share/fonts/truetype/noto/NotoSansCJK.ttc", "NotoSansCJK"),
            ("/usr/share/fonts/opentype/noto-cjk/NotoSansCJK-Regular.ttc", "NotoSansCJK"),
            ("/usr/share/fonts/truetype/noto-cjk/NotoSansCJKjp-Regular.ttf", "NotoSansCJK"),
        ]

        for font_path, font_name in font_paths:
            if os.path.exists(font_path):
                try:
                    # fpdf2 2.x系ではuni=Trueは不要（自動検出）
                    self.add_font(font_name, fname=font_path)
                    self.font_name = font_name
                    logger.info(f"日本語フォント登録成功: {font_path}")
                    return
                except Exception as e:
                    logger.warning(f"フォント登録失敗: {font_path} - {str(e)}")
                    continue

        # フォントが見つからない場合、globで検索
        search_patterns = [
            "/usr/share/fonts/**/NotoSans*CJK*.ttc",
            "/usr/share/fonts/**/NotoSans*CJK*.otf",
            "/usr/share/fonts/**/NotoSans*CJK*.ttf",
            "/usr/share/fonts/**/*Gothic*.ttc",
            "/usr/share/fonts/**/*gothic*.ttc",
        ]
        for pattern in search_patterns:
            try:
                found_fonts = glob.glob(pattern, recursive=True)
                if found_fonts:
                    font_path = found_fonts[0]
                    try:
                        # fpdf2 2.x系ではuni=Trueは不要（自動検出）
                        self.add_font("JapaneseFont", fname=font_path)
                        self.font_name = "JapaneseFont"
                        logger.info(f"日本語フォント登録成功（glob検索）: {font_path}")
                        return
                    except Exception as e:
                        logger.warning(f"フォント登録失敗（glob検索）: {font_path} - {str(e)}")
                        continue
            except Exception as e:
                logger.warning(f"glob検索エラー: {pattern} - {str(e)}")
                continue

        # フォントが見つからない場合
        logger.error("日本語フォントが見つかりません。PDF生成に失敗する可能性があります。")
        self.font_name = None

    def set_japanese_font(self, size=10):
        """日本語フォントを設定"""
        if self.font_name:
            self.set_font(self.font_name, size=size)
        else:
            # フォントがない場合はHelveticaを使用（日本語は表示できない）
            self.set_font("Helvetica", size=size)


class DocumentGenerator:
    def __init__(self):
        """ドキュメント生成の初期化"""
        logger.info("DocumentGenerator初期化完了")

    def _convert_markdown_symbols(self, text: str) -> str:
        """
        Markdown記号を日本語形式に変換
        **テキスト** → 【テキスト】
        """
        # **テキスト** → 【テキスト】
        text = re.sub(r'\*\*(.+?)\*\*', r'【\1】', text)
        return text

    def generate_word(self, content: str, metadata: Dict) -> str:
        """
        Word文書を生成

        Args:
            content: 議事録の本文
            metadata: メタデータ（日付、作成者など）

        Returns:
            生成されたWordファイルのパス
        """
        try:
            logger.info("Word文書の生成を開始")

            # 新しい文書を作成
            doc = Document()

            # タイトル
            title = doc.add_heading('議事録', level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # メタデータテーブル
            doc.add_paragraph()
            table = doc.add_table(rows=4, cols=2)
            table.style = 'Light Grid Accent 1'

            # テーブルの内容
            meta_items = [
                ('作成日', metadata.get('created_date', '')),
                ('作成者', metadata.get('creator', '')),
                ('お客様名', metadata.get('customer_name', '')),
                ('打合せ場所', metadata.get('meeting_place', ''))
            ]

            for i, (label, value) in enumerate(meta_items):
                row = table.rows[i]
                row.cells[0].text = label
                row.cells[1].text = value

                # ラベルセルを太字に
                row.cells[0].paragraphs[0].runs[0].font.bold = True

            # 本文
            doc.add_paragraph()
            doc.add_heading('内容', level=1)

            # Markdown記号を変換
            content = self._convert_markdown_symbols(content)

            # 内容を行ごとに処理
            for line in content.split('\n'):
                line = line.strip()
                if not line:
                    continue

                # セクションヘッダーの判定（##で始まる、または「1. 」〜「9. 」で始まる）
                if line.startswith('##'):
                    heading_text = line.replace('##', '').strip()
                    doc.add_heading(heading_text, level=2)
                elif re.match(r'^[1-9]\.\s', line):
                    # 「1. 打合せ概要」のような形式
                    doc.add_heading(line, level=2)
                # 箇条書きの判定（・、•、-、* で始まる）
                elif line.startswith(('・', '• ', '- ', '* ')):
                    # 箇条書き記号を除去
                    for prefix in ['・', '• ', '- ', '* ']:
                        if line.startswith(prefix):
                            line = line[len(prefix):].strip()
                            break
                    p = doc.add_paragraph(line, style='List Bullet')
                else:
                    # 通常の段落
                    doc.add_paragraph(line)

            # フッター
            doc.add_paragraph()
            footer = doc.add_paragraph(f"作成日時: {datetime.now().strftime('%Y年%m月%d日 %H:%M')}")
            footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            footer_run = footer.runs[0]
            footer_run.font.size = Pt(9)
            footer_run.font.color.rgb = RGBColor(128, 128, 128)

            # 一時ファイルに保存
            output_path = tempfile.mktemp(suffix=".docx")
            doc.save(output_path)

            logger.info(f"Word文書生成完了: {output_path}")
            return output_path

        except Exception as e:
            logger.error(f"Word文書生成エラー: {str(e)}")
            raise

    def generate_pdf(self, content: str, metadata: Dict) -> str:
        """
        PDF文書を生成（fpdf2使用）

        Args:
            content: 議事録の本文
            metadata: メタデータ（日付、作成者など）

        Returns:
            生成されたPDFファイルのパス
        """
        try:
            logger.info("PDF文書の生成を開始")
            logger.info(f"metadata: {metadata}")

            # 一時ファイルパス
            output_path = tempfile.mktemp(suffix=".pdf")

            # PDF作成
            try:
                pdf = JapanesePDF()
                logger.info(f"PDF初期化完了。使用フォント: {pdf.font_name}")
            except Exception as e:
                logger.error(f"PDF初期化エラー: {str(e)}")
                raise

            # ページ設定
            pdf.set_auto_page_break(auto=True, margin=25)
            pdf.set_margins(left=20, top=20, right=20)
            pdf.add_page()

            # 有効なページ幅を計算
            effective_width = pdf.w - pdf.l_margin - pdf.r_margin

            # メタデータ取得
            created_date = str(metadata.get('created_date', '') or '')
            creator = str(metadata.get('creator', '') or '')
            customer_name = str(metadata.get('customer_name', '') or '')
            meeting_place = str(metadata.get('meeting_place', '') or '')

            # 日付フォーマット（ハイフンを除去）
            date_formatted = created_date.replace('-', '')

            # ===== タイトル部分 =====
            # タイトル背景（グラデーション風に2色）
            pdf.set_fill_color(10, 22, 40)  # ダークネイビー
            pdf.rect(pdf.l_margin, pdf.get_y(), effective_width, 22, 'F')

            # アクセントライン
            pdf.set_fill_color(37, 99, 235)  # ブルー
            pdf.rect(pdf.l_margin, pdf.get_y(), 4, 22, 'F')

            # タイトルテキスト（日付_お客様名_議事録）
            title_text = f'{date_formatted}_{customer_name}_議事録'
            pdf.set_text_color(255, 255, 255)
            pdf.set_japanese_font(14)
            pdf.set_xy(pdf.l_margin + 12, pdf.get_y() + 6)
            pdf.cell(effective_width - 12, 10, title_text, align='L')
            pdf.ln(26)

            # ===== メタデータ（モダンカード形式） =====
            # 4列のグリッドレイアウト
            card_width = effective_width / 4
            card_height = 20

            meta_items = [
                ('DATE', created_date),
                ('AUTHOR', creator),
                ('CLIENT', customer_name),
                ('PLACE', meeting_place)
            ]

            # 背景バー
            pdf.set_fill_color(248, 250, 252)  # 薄いグレー背景
            pdf.rect(pdf.l_margin, pdf.get_y(), effective_width, card_height + 4, 'F')

            start_y = pdf.get_y() + 2

            for i, (label, value) in enumerate(meta_items):
                x_pos = pdf.l_margin + (card_width * i) + 4

                # ラベル（小さく、グレー）
                pdf.set_xy(x_pos, start_y)
                pdf.set_japanese_font(7)
                pdf.set_text_color(130, 140, 160)
                pdf.cell(card_width - 8, 4, label, align='L')

                # 値（大きく、ダーク）
                pdf.set_xy(x_pos, start_y + 5)
                pdf.set_japanese_font(9)
                pdf.set_text_color(30, 40, 60)

                # 値が長い場合は切り詰め
                display_value = value if len(value) <= 12 else value[:11] + '...'
                pdf.cell(card_width - 8, 6, display_value, align='L')

                # 区切り線（最後以外）
                if i < len(meta_items) - 1:
                    pdf.set_draw_color(220, 225, 230)
                    line_x = pdf.l_margin + (card_width * (i + 1))
                    pdf.line(line_x, start_y + 2, line_x, start_y + card_height - 4)

            pdf.set_y(start_y + card_height + 6)

            # ===== 区切り線 =====
            pdf.set_draw_color(37, 99, 235)
            pdf.set_line_width(0.5)
            pdf.line(pdf.l_margin, pdf.get_y(), pdf.l_margin + effective_width, pdf.get_y())
            pdf.set_line_width(0.2)
            pdf.ln(10)

            # ===== 本文 =====
            # Markdown記号を変換
            content = self._convert_markdown_symbols(content)

            lines = content.split('\n')
            indent_width = 8
            current_section = None

            for line in lines:
                line = line.strip()

                if not line:
                    pdf.ln(4)
                    continue

                # ## で始まる行は大見出し
                if line.startswith('##'):
                    heading_text = line.replace('##', '').strip()
                    pdf.ln(6)
                    # 見出し背景
                    pdf.set_fill_color(37, 99, 235)  # ブルー
                    pdf.set_text_color(255, 255, 255)
                    pdf.set_japanese_font(11)
                    pdf.cell(effective_width, 10, f'  {heading_text}', fill=True)
                    pdf.ln(12)
                    pdf.set_text_color(0, 0, 0)
                    current_section = heading_text

                # 「1. 」〜「9. 」で始まる行は番号付き見出し
                elif re.match(r'^[1-9]\.\s', line):
                    pdf.ln(6)
                    # 見出し背景
                    pdf.set_fill_color(37, 99, 235)  # ブルー
                    pdf.set_text_color(255, 255, 255)
                    pdf.set_japanese_font(11)
                    pdf.cell(effective_width, 10, f'  {line}', fill=True)
                    pdf.ln(12)
                    pdf.set_text_color(0, 0, 0)
                    current_section = line

                # ・で始まる行は箇条書き
                elif line.startswith('・') or line.startswith('•') or line.startswith('- ') or line.startswith('* '):
                    pdf.set_japanese_font(10)
                    pdf.set_text_color(0, 0, 0)

                    # 箇条書き記号を統一
                    for prefix in ['・', '• ', '- ', '* ']:
                        if line.startswith(prefix):
                            line = line[len(prefix):].strip()
                            break

                    # インデント付きで表示
                    pdf.set_x(pdf.l_margin + indent_width)
                    pdf.set_text_color(37, 99, 235)  # ブルー
                    pdf.cell(5, 7, '●', align='L')
                    pdf.set_text_color(0, 0, 0)
                    pdf.multi_cell(effective_width - indent_width - 5, 7, line)

                # 通常のテキスト
                else:
                    pdf.set_japanese_font(10)
                    pdf.set_text_color(40, 40, 40)
                    pdf.multi_cell(effective_width, 7, line)

            # ===== フッター =====
            pdf.ln(15)
            pdf.set_draw_color(200, 200, 200)
            pdf.line(pdf.l_margin, pdf.get_y(), pdf.l_margin + effective_width, pdf.get_y())
            pdf.ln(5)

            pdf.set_japanese_font(8)
            pdf.set_text_color(128, 128, 128)
            footer_text = f"作成日時: {datetime.now().strftime('%Y年%m月%d日 %H:%M')}"
            pdf.cell(effective_width, 6, footer_text, align='R')

            # PDF保存
            pdf.output(output_path)

            logger.info(f"PDF文書生成完了: {output_path}")
            return output_path

        except Exception as e:
            import traceback
            logger.error(f"PDF文書生成エラー: {str(e)}")
            logger.error(f"スタックトレース: {traceback.format_exc()}")
            raise
